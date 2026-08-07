"""
Report generation for exported search results — PDF, DOCX, XLSX, and
Markdown. Each generator takes a services.report_context.ReportContext
and returns raw bytes; none of this imports Streamlit, so it's usable
from a script, a different frontend, or a future desktop wrapper
exactly as-is.

CSV export lives in services/export.py (it predates this module and
already had its own callers); recommendations_to_rows() from there is
reused here for the XLSX sheet so the two never disagree on columns.
"""

import re
from pathlib import Path
from io import BytesIO

from docx import Document
from openpyxl.utils import get_column_letter
from fpdf import FPDF, XPos, YPos
import pandas as pd

from services.export import recommendations_to_rows
from utils.indexing import format_index_summary

FONT_DIR = Path(__file__).resolve().parent.parent / "assets" / "fonts"

# Zero-width formatting characters (common in Persian/Arabic text) that
# break fpdf2's line-break algorithm even in scripts it otherwise
# renders fine — stripped from every string before it's written.
_ZERO_WIDTH_PATTERN = re.compile("[\u200b\u200c\u200d\u200e\u200f\u202a\u202b\u202c\ufeff]")

# Scripts the bundled DejaVu Sans font either can't shape correctly
# without a proper text-shaping engine (Arabic-family scripts, Thai —
# these need contextual letter joining) or doesn't have glyphs for at
# all (CJK, Hangul). Rather than let fpdf2 raise mid-document — which,
# verified directly, leaves the PDF's internal cursor state corrupted
# for every entry written afterward, not just the failing one — text
# containing these is transliterated to an ASCII-safe fallback before
# it's ever handed to fpdf2. This is a real, disclosed limitation of
# the PDF export specifically; DOCX/XLSX/Markdown have no such issue,
# since they're plain Unicode formats with no font/shaping dependency.
_UNSUPPORTED_SCRIPT_PATTERN = re.compile(
    "[\u0600-\u06FF\u0750-\u077F\u0E00-\u0E7F\u4E00-\u9FFF"
    "\u3040-\u30FF\uAC00-\uD7AF\u3400-\u4DBF\uF900-\uFAFF]"
)


def _pdf_safe_text(text, supported_charset=None):
    """
    Makes text safe to hand to fpdf2 with the bundled DejaVu font.
    Two separate problems handled:
      1. Scripts DejaVu can't shape correctly without a proper text-
         shaping engine (Arabic-family, Thai) — matched by Unicode
         block, since the glyphs exist but joining/shaping doesn't work.
      2. Characters DejaVu has no glyph for at all (CJK, Hangul, emoji,
         some symbols) — checked directly against the font's own glyph
         set, so nothing needs to be individually enumerated by hand.
    Either failure, left unhandled, doesn't just drop the one character —
    verified directly that it corrupts fpdf2's internal cursor state for
    every entry written to the SAME document afterward. So this always
    runs before text reaches multi_cell, never as a try/except around it.
    """
    if not text:
        return text
    text = _ZERO_WIDTH_PATTERN.sub("", str(text))
    if _UNSUPPORTED_SCRIPT_PATTERN.search(text):
        return text.encode("ascii", "replace").decode("ascii")
    if supported_charset is not None:
        text = "".join(
            ch if ord(ch) < 128 or ord(ch) in supported_charset else "?"
            for ch in text
        )
    return text


def _font_charset(pdf):
    """Every codepoint the currently-loaded fonts can actually render."""
    charset = set()
    for font in pdf.fonts.values():
        cmap = getattr(font, "cmap", None)
        if cmap:
            charset |= set(cmap.keys())
    return charset


def _apc_label(result):
    if result["is_free"]:
        return "Free"
    if result["apc_amount"] is not None:
        return f"~${result['apc_amount']:.0f}"
    return "Paid (amount not confirmed in USD)"


def _metadata_lines(context):
    """The small header block requested for every export format."""
    return [
        f"{context.app_name} v{context.app_version}",
        f"Generated: {context.generated_at} UTC",
        "",
        f"Search Strategy: {context.strategy_label}",
        f"Database Sources: {context.database_sources}",
        f"Total Recommendations: {len(context.results)}",
    ]


# ==========================================================
# Markdown — the one most likely to be pasted into an AI
# assistant, so kept plain and clean rather than decorated.
# ==========================================================

def generate_markdown(context) -> bytes:

    lines = []

    lines.extend(_metadata_lines(context))
    lines.append("")
    lines.append("---")
    lines.append("")

    lines.append(f"# {context.title or 'Untitled Search'}")
    lines.append("")

    if context.abstract:
        lines.append("## Abstract")
        lines.append("")
        lines.append(context.abstract)
        lines.append("")

    if context.keywords:
        lines.append(f"**Keywords:** {', '.join(context.keywords)}")
        lines.append("")

    lines.append("## Applied Filters")
    lines.append("")
    for line in context.filters_summary:
        lines.append(f"- {line}")
    lines.append("")

    lines.append("## Recommended Journals")
    lines.append("")

    for i, result in enumerate(context.results, start=1):

        lines.append(f"### {i}. {result['title']}")
        lines.append("")
        lines.append(f"- **Confidence:** {result['confidence']}")
        lines.append(f"- **Indexed in:** {format_index_summary(result['source_details'])}")
        lines.append(f"- **Publisher:** {result['publisher'] or 'Not listed'}")
        lines.append(f"- **Country:** {result['country'] or 'Not listed'}")
        lines.append(f"- **APC:** {_apc_label(result)}")
        lines.append(f"- **Language:** {result['languages'] or 'Not listed'}")
        if result["review_weeks"] is not None:
            lines.append(f"- **Typical review time:** ~{result['review_weeks']} weeks")
        if result.get("explanation"):
            lines.append(f"- **Why this journal:** {result['explanation']}")
        if result["website"]:
            lines.append(f"- **Website:** {result['website']}")
        if result["doaj_url"]:
            lines.append(f"- **DOAJ:** {result['doaj_url']}")
        lines.append("")

    return "\n".join(lines).encode("utf-8")


# ==========================================================
# DOCX
# ==========================================================

def generate_docx(context) -> bytes:

    doc = Document()

    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"{context.app_name} v{context.app_version}")
    meta_run.bold = True
    meta.add_run(f"\nGenerated: {context.generated_at} UTC")
    meta.add_run(f"\nSearch Strategy: {context.strategy_label}")
    meta.add_run(f"\nDatabase Sources: {context.database_sources}")
    meta.add_run(f"\nTotal Recommendations: {len(context.results)}")

    doc.add_heading(context.title or "Untitled Search", level=1)

    if context.abstract:
        doc.add_heading("Abstract", level=2)
        doc.add_paragraph(context.abstract)

    if context.keywords:
        p = doc.add_paragraph()
        p.add_run("Keywords: ").bold = True
        p.add_run(", ".join(context.keywords))

    doc.add_heading("Applied Filters", level=2)
    for line in context.filters_summary:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Recommended Journals", level=2)

    for i, result in enumerate(context.results, start=1):

        doc.add_heading(f"{i}. {result['title']}", level=3)

        p = doc.add_paragraph()
        p.add_run(f"Confidence: {result['confidence']}\n")
        p.add_run(f"Indexed in: {format_index_summary(result['source_details'])}\n")
        p.add_run(f"Publisher: {result['publisher'] or 'Not listed'}\n")
        p.add_run(f"Country: {result['country'] or 'Not listed'}\n")
        p.add_run(f"APC: {_apc_label(result)}\n")
        p.add_run(f"Language: {result['languages'] or 'Not listed'}")
        if result["review_weeks"] is not None:
            p.add_run(f"\nTypical review time: ~{result['review_weeks']} weeks")

        if result.get("explanation"):
            why = doc.add_paragraph()
            why_label = why.add_run("Why this journal: ")
            why_label.italic = True
            why.add_run(result["explanation"])

        links = []
        if result["website"]:
            links.append(f"Website: {result['website']}")
        if result["doaj_url"]:
            links.append(f"DOAJ: {result['doaj_url']}")
        if links:
            doc.add_paragraph(" | ".join(links))

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ==========================================================
# XLSX — reuses recommendations_to_rows from services.export
# so CSV and XLSX never disagree on columns.
# ==========================================================

def generate_xlsx(context) -> bytes:

    info_rows = [
        (context.app_name, f"v{context.app_version}"),
        ("Generated", f"{context.generated_at} UTC"),
        ("Search Strategy", context.strategy_label),
        ("Database Sources", context.database_sources),
        ("Total Recommendations", len(context.results)),
        ("", ""),
        ("Search Title", context.title),
        ("Abstract", context.abstract),
        ("Keywords", ", ".join(context.keywords)),
        ("Applied Filters", "; ".join(context.filters_summary)),
    ]
    info_df = pd.DataFrame(info_rows, columns=["Field", "Value"])

    rec_df = pd.DataFrame(recommendations_to_rows(context.results))

    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        info_df.to_excel(writer, sheet_name="Search Info", index=False, header=False)
        rec_df.to_excel(writer, sheet_name="Recommendations", index=False)

        # Widen columns a bit past their header length so the sheet is
        # readable without the person manually resizing every column.
        sheet = writer.sheets["Recommendations"]
        for i, column in enumerate(rec_df.columns, start=1):
            lengths = rec_df[column].astype(str).str.len()
            max_length = lengths.max()
            width = 12 if pd.isna(max_length) else max(12, min(60, int(max_length) + 2))
            sheet.column_dimensions[get_column_letter(i)].width = width

    return buffer.getvalue()


# ==========================================================
# PDF
# ==========================================================

def generate_pdf(context) -> bytes:

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.add_font("DejaVu", "", str(FONT_DIR / "DejaVuSans.ttf"))
    pdf.add_font("DejaVu", "B", str(FONT_DIR / "DejaVuSans-Bold.ttf"))
    pdf.add_font("DejaVu", "I", str(FONT_DIR / "DejaVuSans-Oblique.ttf"))

    charset = _font_charset(pdf)

    def safe(value):
        return _pdf_safe_text(value, charset)

    def write(text, height=5):
        """
        multi_cell wrapper used for every line in this function.
        Explicit new_x/new_y is load-bearing, not stylistic: verified
        directly that without it, a multi_cell call right after a font
        style switch (e.g. bold heading -> regular body text) can leave
        the cursor in a state where fpdf2 raises "not enough horizontal
        space" on the NEXT call — even for plain ASCII text with no
        Unicode involved at all. Relying on multi_cell's own defaults
        here is not safe with this font/version combination.
        """
        pdf.multi_cell(0, height, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_font("DejaVu", "B", 13)
    write(f"{context.app_name} v{context.app_version}", 7)
    pdf.set_font("DejaVu", "", 9)
    write(
        f"Generated: {context.generated_at} UTC\n"
        f"Search Strategy: {safe(context.strategy_label)}\n"
        f"Database Sources: {context.database_sources}\n"
        f"Total Recommendations: {len(context.results)}"
    )
    pdf.ln(3)

    pdf.set_font("DejaVu", "B", 15)
    write(safe(context.title) or "Untitled Search", 8)
    pdf.ln(1)

    if context.abstract:
        pdf.set_font("DejaVu", "B", 11)
        write("Abstract", 6)
        pdf.set_font("DejaVu", "", 10)
        write(safe(context.abstract))
        pdf.ln(2)

    if context.keywords:
        pdf.set_font("DejaVu", "B", 10)
        write("Keywords: " + safe(", ".join(context.keywords)))
        pdf.ln(1)

    pdf.set_font("DejaVu", "B", 11)
    write("Applied Filters", 6)
    pdf.set_font("DejaVu", "", 10)
    for line in context.filters_summary:
        write(f"- {safe(line)}")
    pdf.ln(3)

    pdf.set_font("DejaVu", "B", 13)
    write("Recommended Journals", 7)
    pdf.ln(1)

    for i, result in enumerate(context.results, start=1):

        pdf.set_font("DejaVu", "B", 11)
        write(f"{i}. {safe(result['title'])}", 6)

        pdf.set_font("DejaVu", "", 9)
        write(
            f"Confidence: {result['confidence']}  |  "
            f"Indexed in: {safe(format_index_summary(result['source_details']))}"
        )
        write(
            f"Publisher: {safe(result['publisher']) or 'Not listed'}  |  "
            f"Country: {safe(result['country']) or 'Not listed'}"
        )
        review_bit = (
            f"  |  Typical review time: ~{result['review_weeks']} weeks"
            if result["review_weeks"] is not None else ""
        )
        write(
            f"APC: {_apc_label(result)}  |  "
            f"Language: {safe(result['languages']) or 'Not listed'}{review_bit}"
        )

        if result.get("explanation"):
            pdf.set_font("DejaVu", "I", 9)
            write(f"Why: {safe(result['explanation'])}")
            pdf.set_font("DejaVu", "", 9)

        links = []
        if result["website"]:
            links.append(f"Website: {result['website']}")
        if result["doaj_url"]:
            links.append(f"DOAJ: {result['doaj_url']}")
        if links:
            write("  |  ".join(links))

        pdf.ln(3)

    return bytes(pdf.output())
